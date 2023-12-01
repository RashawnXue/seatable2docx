import urllib.parse
from datetime import datetime
from io import BytesIO

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from fastapi import FastAPI
from seatable_api import Base
from starlette.responses import StreamingResponse

server_url = ''  # TODO: input your seatable server url
api_token = ''  # TODO: input you api token

base = Base(api_token, server_url)
base.auth()

app = FastAPI()

date = str(datetime.now().date()).replace('-', '.')


def get_data() -> list:
    """
    Get data from seatable
    :return: list that contains data
    """
    rows = base.list_rows("Table1")
    results = []
    for row in rows:
        # TODO: define your custom result here
        results.append(row)
    return results


def gen_doc(data: list) -> Document:
    """
    Generate docx file according to the data
    :param data: the list data
    :return: the docx object
    """
    doc = Document()

    # Set title
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Title")
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = Pt(10.5)
    run.bold = True

    # Set date
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(date)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = Pt(10.5)
    run.bold = True

    for item in data:
        for k, v in item.items():
            if k == 'images':
                # process pictures
                for url in v:
                    try:
                        unquote_url = urllib.parse.unquote('/'.join(url.split('/')[-3:]), encoding='utf-8')
                        download_link = base.get_file_download_link(unquote_url.split('?')[0])
                    except:
                        print("fail: " + url)
                        continue
                    print('success: ' + url)
                    response = requests.get(download_link)
                    binary_img = BytesIO(response.content)
                    doc.add_picture(binary_img, Inches(2))
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(k + '：' + v)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = Pt(10.5)
        doc.add_paragraph()
        doc.add_paragraph()

    return doc


@app.get("/download")
def download_doc():
    print(datetime.now().now())
    doc = gen_doc(get_data())
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    headers = {"Content-Disposition": "attachment; filename=" + date + 'filename.docx'.encode('utf-8').decode(
        'latin-1')}
    return StreamingResponse(content=stream,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers=headers)


if __name__ == '__main__':
    pass
